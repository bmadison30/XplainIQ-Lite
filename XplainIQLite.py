# XplainIQLite.py
# Streamlit app: XplainIQ Channel Readiness Index (Lite)
# Client view: ONLY form (contact fields + 10 sliders). No score/results shown.
# Admin view (?admin=1): show results, radar, download DOCX, optional CSV/XLSX prefill.
# Lead capture to Zapier -> Google Sheets -> CSV fallback
# Manual-review gating: clients submit for review; only Admin (or if gating disabled) can download immediately

import io, os, json, time, base64
from datetime import datetime

import streamlit as st
import pandas as pd
import numpy as np
try:
    import matplotlib.pyplot as plt
    HAS_MPL = True
except Exception:
    HAS_MPL = False
import requests

from docx import Document   # ✅ correct (package is "python-docx")
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# -----------------------------
# Branding / Config (defaults)
# -----------------------------
DEFAULT_BRAND_NAME = "XplainIQ: Channel Readiness Index"
FONT_NAME = "Aptos"
CTA_LINE = "Ready to reach a 90+ Channel Readiness score? Book a full XplainIQ GTM Assessment."
CTA_LINK = "https://calendly.com/"  # <-- set your booking link
FOOTER_NOTE = "© Innovative Networx — XplainIQ™ | Confidential Diagnostic Summary"

# Integrations (env-driven)
ZAPIER_WEBHOOK_URL = os.getenv("ZAPIER_WEBHOOK_URL", "").strip()
GOOGLE_SHEET_ID = os.getenv("GOOGLE_SHEET_ID", "").strip()
GOOGLE_SERVICE_ACCOUNT_JSON = os.getenv("GOOGLE_SERVICE_ACCOUNT_JSON", "").strip()

# Delivery / review gating
APPROVAL_REQUIRED = os.getenv("APPROVAL_REQUIRED", "1").strip().lower() in ("1","true","yes")
# If true, include Base64 DOCX in Zapier payload for Drive/Gmail draft creation
SEND_DOCX_TO_ZAPIER = os.getenv("SEND_DOCX_TO_ZAPIER", "0").strip().lower() in ("1","true","yes")

# -----------------------------
# Pillars / Questions
# -----------------------------
PILLARS = [
    ("A. Channel Strategy & Alignment", ["A1", "A2"]),
    ("B. Partner Program Design",      ["B1", "B2"]),
    ("C. Partner Enablement & Engagement", ["C1", "C2"]),
    ("D. Sales & Operations Integration",  ["D1", "D2"]),
    ("E. Growth Readiness",            ["E1", "E2"]),
]

QUESTIONS = {
    "A1": "Do you have a clearly defined purpose for selling through partners (beyond revenue expansion)?",
    "A2": "Are your target partner types (TSD, VAR, MSP, SI, etc.) well-defined and prioritized?",
    "B1": "Do you have a partner program with tiering, incentives, rules of engagement, or performance criteria?",
    "B2": "Can you clearly articulate what makes your offer unique and profitable for partners?",
    "C1": "Do you provide training, sales playbooks, or co-branded marketing assets?",
    "C2": "How consistently do you communicate and collaborate with active partners?",
    "D1": "Are internal sales/ops aligned to support channel transactions (quoting, order flow, support)?",
    "D2": "Do you track partner pipeline separately with forecast accuracy goals?",
    "E1": "Does senior leadership actively sponsor the channel model?",
    "E2": "Are tools, systems, and staffing sufficient to support 2–3× partner growth?",
}

TIER_BANDS = [
    ("Emerging",   0, 39),
    ("Developing", 40, 59),
    ("Established",60, 79),
    ("Optimized",  80, 100),
]

# -----------------------------
# Helpers
# -----------------------------
def tier_for(score: float) -> str:
    s = round(score)
    for name, lo, hi in TIER_BANDS:
        if lo <= s <= hi:
            return name
    return "Unknown"

def pillar_commentary(pillar_name: str, pscore: float):
    if pscore >= 80:
        return f"{pillar_name} is strong and scalable — keep reinforcing what works."
    if pscore >= 60:
        return f"{pillar_name} shows a solid foundation with room to standardize and scale."
    if pscore >= 40:
        return f"{pillar_name} is emerging — formalize structure, cadence, and measurement."
    return f"{pillar_name} is underdeveloped — prioritize core mechanics and minimum viable structure."

def compute_scores(answers: dict):
    pillar_scores = []
    for pname, qids in PILLARS:
        vals = [int(answers.get(q, 0)) for q in qids]
        if not vals or all(v == 0 for v in vals):
            pscore = 0.0
        else:
            pscore = (sum(vals) / len(vals)) / 5.0 * 100.0
        pillar_scores.append((pname, pscore, dict(zip(qids, vals))))
    overall = sum(p[1] for p in pillar_scores) / len(pillar_scores)
    return pillar_scores, overall

def derive_strengths_gaps(ps):
    sorted_p = sorted(ps, key=lambda x: x[1], reverse=True)
    strengths = [p[0] for p in sorted_p[:2]]
    gaps = [p[0] for p in sorted_p[-3:]]
    return strengths, gaps

def recommend_actions(ps):
    playbook = {
        "A. Channel Strategy & Alignment": "Clarify the partner role by segment and set a 12-month channel thesis with 3 measurable outcomes.",
        "B. Partner Program Design": "Publish a simple one-pager: tiers, incentives, rules of engagement, and co-marketing paths.",
        "C. Partner Enablement & Engagement": "Stand up a 30-60-90 enablement cadence: onboarding kit, monthly enablement call, quarterly MDF campaign.",
        "D. Sales & Operations Integration": "Separate channel pipeline tracking; define lead routing/quoting SLAs; add ‘channel’ to forecast reviews.",
        "E. Growth Readiness": "Baseline partner P&L and capacity; set tooling minimums (PRM/CRM views) and resource triggers for 2–3× growth."
    }
    lows = sorted(ps, key=lambda x: x[1])[:3]
    return [playbook.get(p[0], f"Prioritize foundational improvements in {p[0].lower()} to enable scale.") for p in lows]

def radar_chart(pillar_scores):
    labels = [p[0].split(". ", 1)[1] if ". " in p[0] else p[0] for p in pillar_scores]
    values = [p[1] for p in pillar_scores]
    values = values + values[:1]
    angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
    angles += angles[:1]

    fig = plt.figure()
    ax = plt.subplot(111, polar=True)
    ax.set_theta_offset(np.pi / 2)
    ax.set_theta_direction(-1)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels)
    ax.set_rlabel_position(0)
    ax.set_yticks([20, 40, 60, 80, 100])
    ax.set_ylim(0, 100)
    ax.plot(angles, values, linewidth=2)
    ax.fill(angles, values, alpha=0.1)
    return fig

def image_uploader_return_bytes(label: str):
    up = st.file_uploader(label, type=["png","jpg","jpeg"], accept_multiple_files=False)
    if up:
        return up.read()
    return None

def build_docx(company: str, pillar_scores, overall: float, brand_name: str,
               primary_logo_bytes: bytes|None, partner_logo_bytes: bytes|None,
               tsd_name: str|None) -> bytes:
    doc = Document()

    # Logos row (if provided)
    if primary_logo_bytes or partner_logo_bytes:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        if primary_logo_bytes:
            cell = table.cell(0,0)
            p = cell.paragraphs[0]
            run = p.add_run()
            run.add_picture(io.BytesIO(primary_logo_bytes), width=Inches(1.6))
        if partner_logo_bytes:
            cell = table.cell(0,1)
            p = cell.paragraphs[0]
            run = p.add_run()
            run.add_picture(io.BytesIO(partner_logo_bytes), width=Inches(1.6))
        doc.add_paragraph("")

    # Title
    title = doc.add_paragraph()
    tr = title.add_run(f"{brand_name} — Summary Report")
    tr.bold = True; tr.font.size = Pt(16); tr.font.name = FONT_NAME

    meta = doc.add_paragraph()
    tsd_suffix = f" (Co-branded with {tsd_name})" if tsd_name else ""
    mr = meta.add_run(f"{company}{tsd_suffix} • {datetime.now().strftime('%b %d, %Y')}")
    mr.font.size = Pt(10); mr.font.name = FONT_NAME

    doc.add_paragraph("")

    tier = tier_for(overall)
    p = doc.add_paragraph()
    r = p.add_run(f"Channel Readiness Score: {round(overall)} / 100 — {tier}")
    r.bold = True; r.font.size = Pt(13); r.font.name = FONT_NAME

    doc.add_paragraph("")
    doc.add_paragraph().add_run("Pillar Summary").bold = True
    for pname, pscore, _ in pillar_scores:
        line = doc.add_paragraph()
        run = line.add_run(f"• {pname}: {round(pscore)}")
        run.font.name = FONT_NAME; run.font.size = Pt(11)
        c = doc.add_paragraph(pillar_commentary(pname, pscore))
        c.runs[0].font.name = FONT_NAME; c.runs[0].font.size = Pt(10)

    strengths, gaps = derive_strengths_gaps(pillar_scores)
    doc.add_paragraph("")
    hdr = doc.add_paragraph(); run = hdr.add_run("Top Strengths"); run.bold = True; run.font.name = FONT_NAME
    for s in strengths: doc.add_paragraph(f"• {s}")
    hdr = doc.add_paragraph(); run = hdr.add_run("Opportunities for Improvement"); run.bold = True; run.font.name = FONT_NAME
    for g in gaps: doc.add_paragraph(f"• {g}")

    recs = recommend_actions(pillar_scores)
    hdr = doc.add_paragraph(); run = hdr.add_run("Top 3 Recommendations (Next 90 Days)"); run.bold = True; run.font.name = FONT_NAME
    for rec in recs: doc.add_paragraph(f"• {rec}")

    doc.add_paragraph("")
    cta = doc.add_paragraph()
    rr = cta.add_run(CTA_LINE + " ")
    rr.font.name = FONT_NAME; rr.font.size = Pt(10)
    if CTA_LINK:
        cta.add_run(CTA_LINK).font.name = FONT_NAME

    foot = doc.add_paragraph()
    fr = foot.add_run(FOOTER_NOTE)
    fr.font.size = Pt(8); fr.font.name = FONT_NAME
    foot.alignment = WD_ALIGN_PARAGRAPH.LEFT

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# ----- Lead persistence targets -----
def to_gsheet(row: dict):
    if not GOOGLE_SHEET_ID or not GOOGLE_SERVICE_ACCOUNT_JSON:
        return False, "Google Sheets env vars not set"
    try:
        if GOOGLE_SERVICE_ACCOUNT_JSON.strip().startswith("{"):
            sa_info = json.loads(GOOGLE_SERVICE_ACCOUNT_JSON)
        else:
            with open(GOOGLE_SERVICE_ACCOUNT_JSON, "r") as f:
                sa_info = json.load(f)
        scopes = ["https://www.googleapis.com/auth/spreadsheets"]
        creds = Credentials.from_service_account_info(sa_info, scopes=scopes)
        gc = gspread.authorize(creds)
        sh = gc.open_by_key(GOOGLE_SHEET_ID)
        ws = sh.sheet1
        headers = ws.row_values(1)
        if not headers:
            headers = list(row.keys())
            ws.append_row(headers)
        ordered = [row.get(h, "") for h in headers]
        ws.append_row(ordered)
        return True, "Appended to Google Sheet"
    except Exception as e:
        return False, f"GSheet error: {e}"

def to_zapier(row: dict):
    if not ZAPIER_WEBHOOK_URL:
        return False, "Zapier webhook not set"
    try:
        r = requests.post(ZAPIER_WEBHOOK_URL, json=row, timeout=10)
        if r.status_code // 100 == 2:
            return True, "Sent to Zapier"
        return False, f"Zapier error: {r.status_code} {r.text[:200]}"
    except Exception as e:
        return False, f"Zapier request failed: {e}"

def to_csv(row: dict, path="leads.csv"):
    try:
        exists = os.path.exists(path)
        df = pd.DataFrame([row])
        if exists:
            df.to_csv(path, mode="a", header=False, index=False)
        else:
            df.to_csv(path, index=False)
        return True, f"Saved locally to {path}"
    except Exception as e:
        return False, f"CSV fallback failed: {e}"

def persist_lead(row: dict):
    """Simplified: just save each submission to a local CSV file."""
    try:
        import pandas as pd, os
        path = "leads.csv"
        exists = os.path.exists(path)
        df = pd.DataFrame([row])
        if exists:
            df.to_csv(path, mode="a", header=False, index=False)
        else:
            df.to_csv(path, index=False)
        return True, f"Saved to {path}"
    except Exception as e:
        return False, f"Error saving CSV: {e}"

def get_query_param(key, default=""):
    # Streamlit API migration guard
    try:
        return st.query_params.get(key, default)
    except Exception:
        return st.experimental_get_query_params().get(key, [default])[0]

def prefill_answers_from_query():
    out = {}
    for qid in QUESTIONS.keys():
        try:
            val = int(get_query_param(qid.lower(), ""))
            if 1 <= val <= 5:
                out[qid] = val
        except Exception:
            pass
    return out

# -----------------------------
# UI
# -----------------------------
st.set_page_config(page_title="Channel Readiness Index", page_icon="📊", layout="centered")

# Prefill basics
prefill_company = get_query_param("company", "Unnamed Company")
prefill_name    = get_query_param("name", "")
prefill_email   = get_query_param("email", "")
prefill_role    = get_query_param("role", "")
prefill_phone   = get_query_param("phone", "")
prefill_tsd     = get_query_param("tsd", "")
prefilled_qs    = prefill_answers_from_query()

# Admin mode via URL flag (?admin=1, true, yes)
admin_flag = str(get_query_param("admin", "0")).lower() in ("1", "true", "yes")

st.title("📊 XplainIQ: Channel Readiness Index (Lite)")
if admin_flag:
    st.caption("🔒 Admin mode active (results + download enabled)")
st.write("Please answer the questions below. Your responses will be reviewed and your report will be sent after approval.")

with st.sidebar:
    st.markdown("### Branding")
    brand_name = st.text_input("Brand Name", value=DEFAULT_BRAND_NAME)
    co_brand   = st.toggle("Co-brand with a TSD/Partner?", value=bool(prefill_tsd))
    tsd_name   = st.text_input("TSD/Partner Name", value=prefill_tsd if co_brand else "")
    st.caption("Upload logos (optional)")
    primary_logo_bytes = image_uploader_return_bytes("Primary Logo")
    partner_logo_bytes = image_uploader_return_bytes("Co-brand Logo (TSD)")

    st.markdown("---")
    st.markdown("### Contact")
    name  = st.text_input("Your Name", value=prefill_name)
    email = st.text_input("Work Email", value=prefill_email)
    role  = st.text_input("Title / Role", value=prefill_role)
    phone = st.text_input("Phone (optional)", value=prefill_phone)
    consent = st.checkbox(
        "I consent to Innovative Networx reviewing my responses and contacting me about my results. "
        "I understand my report will be delivered after a manual review.",
        value=True
    )

    st.caption("Submissions are rate-limited per session.")
    if "last_submit_ts" not in st.session_state:
        st.session_state.last_submit_ts = 0.0

    # Admin controls
    if admin_flag:
        st.markdown("---")
        st.markdown("### Admin Tools")
        show_admin_tools = st.toggle("Enable admin upload/prefill?", value=False)
    else:
        show_admin_tools = False

st.markdown("#### Company")
company = st.text_input("Company Name", value=prefill_company)

# Initialize answers (URL prefill supported)
answers = {**prefilled_qs}

# Admin-only: optional CSV/XLSX prefill (hidden for clients)
if show_admin_tools:
    uploaded = st.file_uploader(
        "Admin: Load answers from CSV/Excel (question_id,response,notes)",
        type=["csv", "xlsx", "xls"]
    )
    if uploaded is not None:
        try:
            if uploaded.name.lower().endswith(".csv"):
                df = pd.read_csv(uploaded)
            else:
                df = pd.read_excel(uploaded)
            loaded = 0
            for _, row in df.iterrows():
                qid = str(row.get("question_id", "")).strip()
                if qid in QUESTIONS:
                    try:
                        val = int(row.get("response", 0))
                    except Exception:
                        val = 0
                    val = max(1, min(5, val))
                    answers[qid] = val
                    loaded += 1
            st.success(f"Admin: loaded {loaded} answers from file.")
        except Exception as e:
            st.error(f"Admin: could not parse file: {e}")

# -----------------------------
# QUESTIONS (Client sees only this section)
# -----------------------------
st.markdown("### Questions")
for qid, text in QUESTIONS.items():
    default_val = int(answers.get(qid, 3))  # default to mid-point 3
    answers[qid] = st.slider(f"{qid} — {text}", min_value=1, max_value=5, value=default_val)

# Compute (silently for clients; visible only to admin below)
pillar_scores, overall = compute_scores(answers)
tier = tier_for(overall)

# Generate DOCX bytes (for admin download or Zapier/queue)
if not HAS_DOCX:
    st.warning("DOCX generation is unavailable (python-docx not installed).")
docx_bytes = build_docx(
    company=company,
    pillar_scores=pillar_scores,
    overall=overall,
    brand_name=brand_name,
    primary_logo_bytes=primary_logo_bytes,
    partner_logo_bytes=partner_logo_bytes if co_brand else None,
    tsd_name=tsd_name if co_brand else None,
)
stamp = datetime.now().strftime("%Y%m%d_%H%M")
fname = f"{company.replace(' ','')}_ChannelReadiness_{stamp}.docx"

# -----------------------------
# ADMIN-ONLY RESULTS & DOWNLOAD
# -----------------------------
if admin_flag:
    st.markdown("---")
    st.subheader("Admin: Results Preview")
    c1, c2 = st.columns(2)
    with c1:
        st.metric(label="Channel Readiness Score", value=f"{round(overall)} / 100")
    with c2:
        st.metric(label="Maturity Tier", value=tier)

    st.subheader("Readiness Radar")
    if HAS_MPL:
        fig = radar_chart(pillar_scores)
        st.pyplot(fig, use_container_width=True)
    else:
        st.info("Radar chart unavailable (matplotlib not installed).")
    for pname, pscore, detail in pillar_scores:
        with st.expander(f"{pname}: {round(pscore)}"):
            st.write(pillar_commentary(pname, pscore))
            st.write(f"Q detail: {detail}")

    st.subheader("Readiness Radar")
    fig = radar_chart(pillar_scores)
    st.pyplot(fig, use_container_width=True)

    # Admin can download immediately
    if admin_flag or not APPROVAL_REQUIRED:
        st.download_button(
            "📄 Download 1-Page DOCX",
            data=docx_bytes,
            file_name=fname,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

# -----------------------------
# SUBMIT (Lead capture, no client-facing results)
# -----------------------------
st.markdown("---")
submit_label = "Submit for Review" if not admin_flag else "Submit & Record"
if st.button(submit_label):
    now = time.time()
    if now - st.session_state.last_submit_ts < 60:
        st.warning("Please wait a minute before submitting again.")
    elif not consent:
        st.error("Please provide consent to proceed.")
    elif not email or "@" not in email:
        st.error("Please enter a valid work email.")
    else:
        st.session_state.last_submit_ts = now
        payload = {
            "ts": datetime.utcnow().isoformat() + "Z",
            "brand_name": brand_name,
            "tsd_cobrand": tsd_name if co_brand else "",
            "company": company,
            "name": name,
            "email": email,
            "role": role,
            "phone": phone,
            "score_overall": round(overall),
            "tier": tier,
            "pillar_scores": {p[0]: round(p[1]) for p in pillar_scores},
            "answers": answers,
            # Review gating fields
            "status": "Pending Review" if not admin_flag else "Submitted by Admin",
            "approval_required": bool(APPROVAL_REQUIRED),
        }

        # Optional: attach DOCX for Zapier (Base64) so your Zap can store/attach the file
        if SEND_DOCX_TO_ZAPIER and ZAPIER_WEBHOOK_URL:
            payload["docx_filename"] = fname
            payload["docx_b64"] = base64.b64encode(docx_bytes).decode("utf-8")

        ok, msg = persist_lead(payload)
        if ok:
            if admin_flag:
                st.success("Recorded. (Admin mode)")
            else:
                st.success("Submitted for review. We’ll email your report after approval.")
        else:
            st.error(f"Submission error: {msg}")

st.caption("Powered by XplainIQ™ • Engineering Predictable Go-To-Market Outcomes.")





